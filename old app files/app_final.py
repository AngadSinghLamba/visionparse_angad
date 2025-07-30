# app.py  ── VisionParse 5.6 ----------------------------------------------
# PDF/Word/Excel → MD, HTML, TXT, JSON, XML + Assets → Nested ZIP download
# Robust PDF handling via temp files + PyPDF2 pre‑check
# ----------------------------------------------------------------------------

import streamlit as st
import warnings, zipfile, base64, time, json, os, tempfile
from pathlib import Path
from dataclasses import dataclass, field
from typing import List
import io as _io
import pandas as pd
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
from docling.datamodel.base_models import DocumentStream
from docling_core.types.doc import ImageRefMode, PictureItem, TableItem
from docling.datamodel.pipeline_options import (
    PdfPipelineOptions, TableFormerMode, EasyOcrOptions
)
from docling.document_converter import (
    DocumentConverter, PdfFormatOption, WordFormatOption
)
from docling.pipeline.standard_pdf_pipeline import StandardPdfPipeline
from docling.pipeline.simple_pipeline      import SimplePipeline
from docling.backend.pypdfium2_backend      import PyPdfiumDocumentBackend
from docling.datamodel.base_models          import InputFormat

warnings.filterwarnings(
    "ignore",
    message="'pin_memory' argument is set as true but not supported on MPS"
)

@dataclass
class VisionParseConfig:
    SUPPORTED_TYPES: List[str] = field(default_factory=lambda: ["pdf", "docx", "xlsx"])
    DEFAULT_IMAGE_SCALE: float = 3.0
    OUTPUT_DIR: Path = Path("artifacts")

def init_session():
    st.session_state.setdefault("files", [])

@st.cache_resource(show_spinner="🔧 Loading Docling & EasyOCR…")
def get_converter(easy_ocr: bool, extract_images: bool, extract_tables: bool, image_scale: float):
    pdf_opts = PdfPipelineOptions(
        do_ocr=easy_ocr,
        ocr_options=EasyOcrOptions(force_full_page_ocr=True, lang=["en"]) if easy_ocr else None,
        do_table_structure=extract_tables,
        table_structure_options=dict(mode=TableFormerMode.ACCURATE, do_cell_matching=False) if extract_tables else None,
        generate_page_images=extract_images,
        generate_picture_images=extract_images,
        images_scale=image_scale,
    )
    return DocumentConverter(
        allowed_formats=[InputFormat.PDF, InputFormat.DOCX],
        format_options={
            InputFormat.PDF: PdfFormatOption(
                pipeline_cls=StandardPdfPipeline,
                backend=PyPdfiumDocumentBackend,
                pipeline_options=pdf_opts
            ),
            InputFormat.DOCX: WordFormatOption(pipeline_cls=SimplePipeline)
        }
    )

def make_download_link(buf: bytes, name="visionparse_output.zip"):
    b64 = base64.b64encode(buf).decode()
    return f'<a href="data:application/zip;base64,{b64}" download="{name}">⬇️ Download ZIP</a>'

def export_pdf_tables(res_document, base: Path, extract_tables: bool):
    paths = []
    if extract_tables and hasattr(res_document, 'tables'):
        for i, tbl in enumerate(res_document.tables, start=1):
            df = tbl.export_to_dataframe()
            p  = base / f"{base.name}_table_{i}.csv"
            df.to_csv(p, index=False)
            paths.append(p)
    return paths

def export_docx_assets(data: bytes, base: Path, extract_images: bool, extract_tables: bool):
    imgs, tbls = [], []
    doc = DocxDocument(_io.BytesIO(data))
    if extract_images:
        for i, rel in enumerate(doc.part._rels.values(), start=1):
            if "image" in rel.target_ref:
                blob = rel.target_part.blob
                p = base / f"{base.name}_image_{i}.png"
                p.write_bytes(blob)
                imgs.append(p)
    if extract_tables:
        for j, table in enumerate(doc.tables, start=1):
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            df = pd.DataFrame(rows)
            p  = base / f"{base.name}_table_{j}.csv"
            df.to_csv(p, index=False)
            tbls.append(p)
    return imgs, tbls

def generate_annotated_txt(res, base: Path):
    lines, tcount, icount = [], 0, 0
    img_files = sorted(base.glob("*.png"))
    for el, _ in res.document.iterate_items():
        lvl = getattr(el, "level", None)
        txt = getattr(el, "text", "").strip()
        if isinstance(lvl, int) and txt:
            lines.append(f"{'#'*lvl} {txt}")
            continue
        if isinstance(el, TableItem):
            tcount += 1
            df = el.export_to_dataframe()
            lines.append(f"[Table {tcount}]")
            lines.append(df.to_string(index=False))
            continue
        if isinstance(el, PictureItem):
            fname = img_files[icount].name if icount < len(img_files) else f"{base.name}_image_{icount+1}.png"
            icount += 1
            lines.append(f"[Image: {fname}]")
            continue
        if txt:
            lines.append(txt)
    txt_p = base / f"{base.name}.txt"
    txt_p.write_text("\n\n".join(lines), encoding="utf-8")
    return txt_p

def main():
    cfg = VisionParseConfig()
    st.set_page_config("VisionParse 5.6", "📄", "wide")
    st.title("📄 VisionParse 5.6 – Robust PDF/Word/Excel → MD, HTML, TXT, JSON, XML + Assets")

    init_session()

    easy_ocr       = st.sidebar.checkbox("Enable EasyOCR (PDF)", True)
    extract_tables = st.sidebar.checkbox("Extract Tables", True)
    extract_images = st.sidebar.checkbox("Extract Images", True)
    export_xml     = st.sidebar.checkbox("Export DocTag XML", False)
    image_scale    = st.sidebar.slider("Image Scale (PDF)", 1.0, 4.0, cfg.DEFAULT_IMAGE_SCALE, 0.5)
    max_pages      = st.sidebar.number_input("Max pages (PDF)", 1, 500, 10)

    files = st.file_uploader(
        "Upload PDF, Word (.docx), Excel (.xlsx)",
        type=cfg.SUPPORTED_TYPES, accept_multiple_files=True
    )
    if files:
        st.session_state.files = files

    if st.button("🚀 Convert All", disabled=not st.session_state.files):
        converter = get_converter(easy_ocr, extract_images, extract_tables, image_scale)
        out_root  = cfg.OUTPUT_DIR
        out_root.mkdir(exist_ok=True)

        for f in st.session_state.files:
            base = out_root / Path(f.name).stem
            base.mkdir(exist_ok=True)
            data = f.read()
            ext  = Path(f.name).suffix.lower().strip(".")

            # EXCEL handling
            if ext == "xlsx":
                if extract_tables:
                    sheets = pd.read_excel(_io.BytesIO(data), sheet_name=None)
                    for sh, df in sheets.items():
                        pfx = f"{base.name}_{sh}"
                        (base/f"{pfx}_table.csv").write_text(df.to_csv(index=False), encoding="utf-8")
                        (base/f"{pfx}.md").write_text(df.to_markdown(index=False), encoding="utf-8")
                        (base/f"{pfx}.html").write_text(df.to_html(index=False), encoding="utf-8")
                        (base/f"{pfx}.txt").write_text(df.to_string(index=False), encoding="utf-8")
                continue

            # WORD handling
            if ext == "docx":
                imgs, tbls = export_docx_assets(data, base, extract_images, extract_tables)
                src = DocumentStream(name=f.name, stream=_io.BytesIO(data))
                res = converter.convert(src, max_num_pages=max_pages)

                # Markdown, HTML, JSON, XML
                md_p, html_p = base / f"{base.name}.md", base / f"{base.name}.html"
                json_p = base / f"{base.name}.json"
                res.document.save_as_markdown(md_p, image_mode=ImageRefMode.REFERENCED)
                res.document.save_as_html  (html_p, image_mode=ImageRefMode.REFERENCED)
                json_p.write_text(json.dumps(res.document.export_to_dict(), ensure_ascii=False), encoding="utf-8")
                if export_xml:
                    xml_p = base / f"{base.name}.xml"
                    res.document.save_as_xml(xml_p)

                generate_annotated_txt(res, base)
                continue

            # PDF handling with pre‑check & temp file
            if ext == "pdf":
                # sanity-check
                try:
                    PdfReader(_io.BytesIO(data))
                except Exception:
                    st.error(f"🚨 Invalid PDF file: {f.name} – skipping.")
                    continue

                # write temp file for pypdfium2 backend
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                    tmp.write(data)
                    tmp_path = tmp.name

                t0 = time.time()
                try:
                    # pass path string to converter
                    res = converter.convert(tmp_path, max_num_pages=max_pages)
                except Exception as e:
                    st.error(f"Error during conversion of {f.name}: {e}")
                    os.unlink(tmp_path)
                    continue
                os.unlink(tmp_path)
                st.sidebar.write(f"{f.name}: {time.time()-t0:.1f}s")

                md_p, html_p = base / f"{base.name}.md", base / f"{base.name}.html"
                json_p = base / f"{base.name}.json"
                res.document.save_as_markdown(
                    md_p,
                    image_mode=ImageRefMode.REFERENCED if extract_images else "text"
                )
                res.document.save_as_html  (
                    html_p,
                    image_mode=ImageRefMode.REFERENCED if extract_images else "text"
                )
                json_p.write_text(json.dumps(res.document.export_to_dict(), ensure_ascii=False), encoding="utf-8")
                if export_xml:
                    xml_p = base / f"{base.name}.xml"
                    res.document.save_as_xml(xml_p)

                export_pdf_tables(res.document, base, extract_tables)
                generate_annotated_txt(res, base)
                continue

        # Build nested ZIP
        zip_buf = _io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for doc_dir in out_root.iterdir():
                if not doc_dir.is_dir(): continue
                name = doc_dir.name
                for file_path in doc_dir.iterdir():
                    ext = file_path.suffix.lower()
                    if ext == ".md":
                        arc = f"{name}/md/{file_path.name}"
                    elif ext == ".html":
                        arc = f"{name}/html/{file_path.name}"
                    elif ext == ".txt":
                        arc = f"{name}/txt/{file_path.name}"
                    elif ext == ".json":
                        arc = f"{name}/json/{file_path.name}"
                    elif ext == ".xml":
                        arc = f"{name}/xml/{file_path.name}"
                    elif ext == ".csv":
                        arc = f"{name}/assets/tables/{file_path.name}"
                    elif ext in [".png", ".jpg", ".jpeg"]:
                        arc = f"{name}/assets/images/{file_path.name}"
                    else:
                        arc = f"{name}/{file_path.name}"
                    zf.write(file_path, arc)

        st.success("✅ All done!")
        st.markdown(make_download_link(zip_buf.getvalue()), unsafe_allow_html=True)

    if st.button("🗑️ Clear session"):
        st.session_state.files.clear()
        st.experimental_rerun()

if __name__ == "__main__":
    main()
